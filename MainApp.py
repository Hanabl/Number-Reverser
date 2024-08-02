from docx import Document

# def reverse_text(text):
#     return text[::-1]



fileAddress = 'primaryFile.docx'
doc = Document(fileAddress)

new_doc = Document()



for para in doc.paragraphs:
        print(para.text)
        

# new_doc.save('reversed_output.docx')

print('file saved')